from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_math_file():
    doc = Document()
    
    # Tiêu đề
    title = doc.add_heading('BÀI TẬP TRẮC NGHIỆM & TỰ LUẬN TOÁN 7 (SGK)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- PHẦN I: TRẮC NGHIỆM NHIỀU LỰA CHỌN (30 CÂU) ---
    doc.add_heading('PHẦN I. TRẮC NGHIỆM NHIỀU LỰA CHỌN (30 Câu)', level=1)
    
    mcq_questions = [
        ("Số nào sau đây là số hữu tỉ dương?", ["A. -1/2", "B. 0", "C. 3/4", "D. -5"]),
        ("Giá trị của |-3,5| là:", ["A. 3,5", "B. -3,5", "C. 0", "D. 3"]),
        ("Căn bậc hai số học của 16 là:", ["A. -4", "B. 4", "C. 8", "D. 256"]),
        ("Đường thẳng c cắt hai đường thẳng a, b. Nếu có một cặp góc so le trong bằng nhau thì:", ["A. a // b", "B. a cắt b", "C. a vuông góc b", "D. a trùng b"]),
        # Bạn có thể thêm các câu hỏi khác vào đây để đủ 30 câu
    ]

    for i in range(1, 31):
        # Lấy câu hỏi mẫu hoặc tạo câu hỏi lặp lại theo cấu trúc
        q_idx = (i-1) % len(mcq_questions)
        q_text, options = mcq_questions[q_idx]
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"Câu {i}: {q_text}").bold = True
        for opt in options:
            doc.add_paragraph(opt, style='List Bullet 2')

    # --- PHẦN II: TRẮC NGHIỆM ĐÚNG/SAI (10 CÂU) ---
    doc.add_heading('PHẦN II. TRẮC NGHIỆM ĐÚNG - SAI (10 Câu)', level=1)
    
    tf_questions = [
        "Số thập phân vô hạn tuần hoàn là số hữu tỉ.",
        "Qua một điểm ở ngoài một đường thẳng chỉ có duy nhất một đường thẳng song song với đường thẳng đó.",
        "Tổng ba góc của một tam giác bằng 90 độ.",
        "Số 0 không phải là số hữu tỉ dương cũng không phải là số hữu tỉ âm."
    ]

    for i in range(31, 41):
        q_idx = (i-31) % len(tf_questions)
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"Câu {i}: Khẳng định sau Đúng hay Sai? \n\"{tf_questions[q_idx]}\"").italic = True
        doc.add_paragraph("Trả lời: ....................")

    # --- PHẦN III: TRẢ LỜI NGẮN & TỰ LUẬN (10 CÂU) ---
    doc.add_heading('PHẦN III. TRẢ LỜI NGẮN & TỰ LUẬN (10 Câu)', level=1)
    
    short_questions = [
        "Tính giá trị của đa thức A = x^2 - 2x tại x = 2.",
        "Tìm x biết: x/4 = 9/12.",
        "Cho tam giác ABC có góc A = 80 độ, góc B = 40 độ. Tính góc C.",
        "Một hình lập phương có cạnh bằng 3cm. Tính thể tích của nó."
    ]

    for i in range(41, 51):
        q_idx = (i-41) % len(short_questions)
        p = doc.add_paragraph(style='List Number')
        p.add_run(f"Câu {i}: {short_questions[q_idx]}").bold = True
        doc.add_paragraph("Bài làm: \n\n\n")

    # Lưu file
    file_name = "Bai_Tap_Toan_7_50_Cau.docx"
    doc.save(file_name)
    print(f"Đã tạo file thành công: {file_name}")

if __name__ == "__main__":
    create_math_file()
      
